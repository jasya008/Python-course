from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys


def create_letter(text_lines):
    document = Document()

    # Greeting
    greeting_paragraph = document.add_paragraph(text_lines[0])
    greeting_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    greeting_run = greeting_paragraph.runs[0]
    greeting_run.font.name = 'Arial'
    greeting_run.font.size = Pt(11)
    greeting_run.italic = True

    # Body paragraphs
    for line in text_lines[1:-2]:
        body_paragraph = document.add_paragraph(line)
        body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_paragraph.paragraph_format.first_line_indent = Inches(0.5)

        body_run = body_paragraph.runs[0]
        body_run.font.name = 'Times New Roman'
        body_run.font.size = Pt(12)

        body_run_element = body_run._element
        body_run_element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Closing
    for line in text_lines[-2:]:
        closing_paragraph = document.add_paragraph(line)
        closing_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        closing_run = closing_paragraph.runs[0]
        closing_run.font.name = 'Times New Roman'
        closing_run.font.size = Pt(12)

        if line == text_lines[-1]:
            closing_run.bold = True

        closing_run_element = closing_run._element
        closing_run_element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Save the document
    document.save('letter.docx')

text_lines = []

create_letter(text_lines)
